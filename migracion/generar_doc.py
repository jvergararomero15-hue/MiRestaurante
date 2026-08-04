from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ========== ESTILOS ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6a)
    hs.font.name = 'Calibri'

# ========== PORTADA ==========
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SISTEMA DE GESTIÓN\nRESTAURANTE')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0xf3, 0xb3, 0x3d)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Documento de Presentación del Proyecto')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6a)

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Fecha: {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(12)

teclist = doc.add_paragraph()
teclist.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = teclist.add_run('Stack: Java 21 · Spring Boot 3.5 · React 19 · PostgreSQL · Docker')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ========== ÍNDICE (manual) ==========
doc.add_heading('Índice', level=1)
items = [
    '1. Resumen Ejecutivo',
    '2. Arquitectura del Sistema',
    '3. Tecnologías Utilizadas',
    '4. Modelo de Datos',
    '5. Módulos del Sistema',
    '6. Seguridad y Autenticación',
    '7. API REST',
    '8. Panel Administrador',
    '9. Despliegue y Docker',
    '10. Pruebas y Calidad',
]
for item in items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ========== 1. RESUMEN EJECUTIVO ==========
doc.add_heading('1. Resumen Ejecutivo', level=1)
doc.add_paragraph(
    'El Sistema de Gestión para Restaurante es una aplicación web full-stack diseñada para '
    'administrar las operaciones diarias de un restaurante. Permite gestionar mesas, menú, '
    'reservas, pedidos y cobros, con roles diferenciados para clientes y administradores.'
)
doc.add_paragraph(
    'El sistema cuenta con autenticación JWT, panel de administración con dashboard de '
    'estadísticas en tiempo real, y soporte para contenedores Docker para facilitar el despliegue.'
)

# ========== 2. ARQUITECTURA ==========
doc.add_heading('2. Arquitectura del Sistema', level=1)
doc.add_paragraph('El sistema sigue una arquitectura de tres capas:')

arq = [
    ('Frontend (React 19 + Vite 8)', 
     'Interfaz de usuario moderna y responsiva. Incluye React Router para navegación '
     'del lado del cliente, Axios para comunicación HTTP, y estilos CSS personalizados. '
     'Se sirve mediante nginx en producción.'),
    ('Backend (Spring Boot 3.5 + Java 21)', 
     'API REST con Spring Data JPA para persistencia, Spring Security con JWT para '
     'autenticación, y validación de inputs con Jakarta Validation. Arquitectura en '
     'capas: Controller → Service → Repository.'),
    ('Base de Datos (PostgreSQL)', 
     'Base de datos relacional con tablas normalizadas para usuarios, mesas, platos, '
     'pedidos, reservas, clientes y detalles de pedidos. Migraciones automáticas mediante DDL Auto.'),
]

for title, desc in arq:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(desc)

# Diagrama ASCII-like
doc.add_heading('Diagrama de Arquitectura', level=2)
diagram = """  [ Navegador ]
        |
  [ nginx (proxy reverse) ]
        |
  -------------------------
  |    React App (Vite)   |
  |  Puerto 5173 (dev)     |
  |  Puerto 80 (prod)      |
  -------------------------
        |  HTTP / JSON
  -------------------------
  |  Spring Boot Backend  |
  |  Puerto 8080           |
  |  JWT Auth              |
  -------------------------
        |  JDBC
  -------------------------
  |  PostgreSQL           |
  |  Puerto 5432           |
  -------------------------"""

p = doc.add_paragraph()
run = p.add_run(diagram)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# ========== 3. TECNOLOGÍAS ==========
doc.add_heading('3. Tecnologías Utilizadas', level=1)

table = doc.add_table(rows=10, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Capa', 'Tecnología', 'Versión']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

techs = [
    ('Frontend', 'React', '19.2'),
    ('Frontend', 'Vite', '8.0'),
    ('Frontend', 'React Router', '7.14'),
    ('Frontend', 'Axios', '1.18'),
    ('Backend', 'Java', '21'),
    ('Backend', 'Spring Boot', '3.5.6'),
    ('Backend', 'Spring Security + JWT', '0.12.7'),
    ('Base Datos', 'PostgreSQL', '16+'),
    ('Infra', 'Docker / Docker Compose', 'Última'),
]

for row_idx, (layer, tech, ver) in enumerate(techs, 1):
    table.rows[row_idx].cells[0].text = layer
    table.rows[row_idx].cells[1].text = tech
    table.rows[row_idx].cells[2].text = ver

doc.add_paragraph('')

# ========== 4. MODELO DE DATOS ==========
doc.add_heading('4. Modelo de Datos', level=1)
doc.add_paragraph('El sistema cuenta con 7 entidades principales:')

entities = [
    ('Usuario', 'id_usuario, nombre, apellido, correo, password, rol (USER/ADMIN), estado, fecha_registro'),
    ('Cliente', 'id_cliente, nombre, apellido, cedula, telefono, correo, direccion, fecha_registro, estado'),
    ('Mesa', 'id_mesa, numero, capacidad, estado (Libre/Ocupada/Reservada), reservado_por'),
    ('Plato', 'id_plato, nombre, descripcion, precio, imagen, categoria, disponible'),
    ('Reserva', 'id_reserva, fecha, hora, cantidad_personas, estado (Activa/Cancelada), FK → Mesa, FK → Cliente'),
    ('Pedido', 'id_pedido, fecha, total, estado (Activo/Cobrado), FK → Mesa, FK → Cliente, FK → Usuario'),
    ('DetallePedido', 'id_detalle, cantidad, precio_unitario, subtotal, FK → Pedido, FK → Plato'),
]

for name, fields in entities:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(fields)

doc.add_page_break()

# ========== 5. MÓDULOS ==========
doc.add_heading('5. Módulos del Sistema', level=1)

modulos = [
    ('Autenticación y Usuarios', 
     'Registro e inicio de sesión con JWT. Roles: ADMIN (control total) y USER (cliente). '
     'Las contraseñas se almacenan cifradas con BCrypt.'),
    ('Gestión de Menú', 
     'CRUD completo de platos con nombre, descripción, precio, imagen (emoji) y categoría. '
     'El administrador puede agregar, editar y eliminar platos.'),
    ('Gestión de Mesas', 
     'Control de estado de mesas: Libre, Ocupada, Reservada. Permite agregar/eliminar '
     'productos al consumo de cada mesa y calcular total automáticamente.'),
    ('Reservas', 
     'Los clientes pueden reservar mesas seleccionando fecha, hora y número de personas. '
     'El sistema valida que no haya conflictos de horario para la misma mesa.'),
    ('Pedidos y Cobro', 
     'Al cobrar una mesa, se genera automáticamente un pedido en la base de datos con el '
     'total de los consumos. Solo los administradores pueden ejecutar el cobro.'),
    ('Panel de Administración', 
     'Dashboard con estadísticas en tiempo real: mesas libres/ocupadas/reservadas, '
     'reservas del día, pedidos del día, ventas totales y del día. CRUD de clientes, '
     'mesas, platos y visualización de historial de ventas y reservas.'),
]

for title, desc in modulos:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph(desc)
    p2.paragraph_format.space_after = Pt(10)

# ========== 6. SEGURIDAD ==========
doc.add_heading('6. Seguridad y Autenticación', level=1)
doc.add_paragraph('La seguridad se implementa en múltiples capas:')

seg = [
    'JWT (JSON Web Token): Autenticación stateless. El token se genera al login y se envía en el header Authorization de cada petición.',
    'BCrypt: Las contraseñas se almacenan cifradas. Nunca viajan en texto plano.',
    'Roles: Separación clara entre USER y ADMIN. Rutas del panel admin protegidas.',
    'CORS: Configuración de orígenes permitidos para evitar accesos no autorizados.',
    'Validación: Todos los inputs se validan con Jakarta Validation (@NotBlank, @Email, @Positive, etc.).',
    'Manejo global de errores: @ControllerAdvice captura excepciones y responde con códigos HTTP apropiados (400, 404, 409, 500).',
]

for s in seg:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ========== 7. API REST ==========
doc.add_heading('7. API REST - Endpoints Principales', level=1)

table = doc.add_table(rows=14, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Método', 'Endpoint', 'Descripción']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows = [
    ('POST', '/auth/registro', 'Registro de nuevo usuario'),
    ('POST', '/auth/login', 'Inicio de sesión, devuelve JWT'),
    ('GET', '/platos', 'Listar todos los platos'),
    ('GET', '/mesas', 'Listar todas las mesas'),
    ('PUT', '/mesas/{id}', 'Actualizar estado de una mesa'),
    ('GET', '/reservas', 'Listar todas las reservas'),
    ('POST', '/reservas', 'Crear una nueva reserva'),
    ('POST', '/pedidos', 'Crear un pedido (al cobrar)'),
    ('GET', '/admin/dashboard', 'Estadísticas del dashboard'),
    ('GET', '/admin/ventas', 'Historial de ventas por fecha'),
    ('GET', '/clientes', 'Listar clientes'),
    ('POST', '/clientes', 'Crear cliente'),
    ('DELETE', '/clientes/{id}', 'Eliminar cliente'),
]

for row_idx, (method, endpoint, desc) in enumerate(rows, 1):
    table.rows[row_idx].cells[0].text = method
    table.rows[row_idx].cells[1].text = endpoint
    table.rows[row_idx].cells[2].text = desc

# ========== 8. PANEL ADMIN ==========
doc.add_heading('8. Panel de Administración', level=1)
doc.add_paragraph(
    'El panel admin es accesible en /admin para usuarios con rol ADMIN. '
    'Incluye las siguientes secciones:'
)

sections = [
    ('Dashboard', 'Tarjetas con estadísticas: mesas libres/ocupadas/reservadas, reservas totales y del día, pedidos totales y del día, ventas totales y del día.'),
    ('Reservas', 'Historial completo de reservas con filtro por fecha y opción de eliminar.'),
    ('Ventas', 'Historial de pedidos cobrados con filtro por fecha y total acumulado del día.'),
    ('Menú', 'CRUD de platos: agregar, editar y eliminar.'),
    ('Clientes', 'Listado de todos los clientes registrados con opción de eliminar.'),
    ('Mesas', 'Listado de mesas con opción de liberar o eliminar.'),
]

for title, desc in sections:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== 9. DESPLIEGUE ==========
doc.add_heading('9. Despliegue con Docker', level=1)
doc.add_paragraph(
    'El sistema está completamente containerizado para facilitar el despliegue '
    'en cualquier entorno. Incluye:'
)

docker_items = [
    'Dockerfile para el backend (multi-stage: build Maven → JRE 21)',
    'Dockerfile para el frontend (multi-stage: build Node → nginx)',
    'docker-compose.yml con 3 servicios: PostgreSQL, backend, frontend',
    'Nginx configurado como proxy reverso para las rutas API',
    'Healthcheck de base de datos para asegurar el orden de inicio',
    'Volumen persistente para los datos de PostgreSQL',
]

for item in docker_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_heading('Comandos principales', level=2)
commands = [
    ('Desarrollo', 'mvn spring-boot:run (backend) | npm run dev (frontend)'),
    ('Construir JAR', 'mvn clean package -DskipTests'),
    ('Docker', 'docker compose up --build'),
    ('Tests', 'mvn test'),
]
for name, cmd in commands:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}: ')
    run.bold = True
    run = p.add_run(cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

# ========== 10. PRUEBAS ==========
doc.add_heading('10. Pruebas y Calidad', level=1)
doc.add_paragraph(
    'El sistema incluye pruebas unitarias con Mockito para los servicios principales. '
    'Se implementaron 15 tests automatizados que cubren:'
)

tests = [
    'MesaService (4 tests): listar, buscar, guardar, eliminar',
    'PedidoService (3 tests): listar, guardar, buscar',
    'PlatoService (3 tests): listar, guardar, eliminar',
    'ClienteService (2 tests): listar, guardar',
    'ReservaService (2 tests): listar, guardar',
    'ContextLoad (1 test): Verifica que la aplicación arranca',
]
for t in tests:
    doc.add_paragraph(t, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Resultado: 15 tests, 0 fallos')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x4c, 0xaf, 0x50)

# ========== FINAL ==========
doc.add_page_break()
doc.add_heading('Contacto', level=1)
doc.add_paragraph('Este proyecto fue desarrollado como sistema de gestión integral para restaurante.')
doc.add_paragraph('Para más información o soporte técnico, contactar al equipo de desarrollo.')

# ========== GUARDAR ==========
output_path = 'C:/Users/joseb/Documents/ProyectosVSC/Restaurante-Full/Presentacion_Proyecto_Restaurante.docx'
doc.save(output_path)
print(f'Documento creado: {output_path}')
